import os
import re
import pprint
import json
import pandas as pd
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment
import sys
from docx import Document
from gherkin.parser import Parser
from FeatureErrorHandler import FeatureErrorHandler
# Alterado para usar a versão headless
from planilha_regressao_json_jenkins import gerar_planilha_regressao, converte_json
#
def validar_feature_estrutura(dados, caminho_relativo,nome_arquivo, erros):
    feature = dados
    children = feature.get("children", [])
    for bloco in children:
        scenario = bloco.get("scenario", {})
        keyword = scenario.get("keyword", "")
        examples = scenario.get("examples", [])
        if keyword == "Scenario Outline" and not examples:
            erros.append({
                "Arquivo": caminho_relativo,
                "Tipo do erro": "Estrutura",
                "Descrição do erro": "'Scenario Outline' encontrado sem 'Examples'."
            })
        elif keyword == "Scenario" and examples:
            erros.append({
                "Arquivo": caminho_relativo,
                "Tipo do erro": "Estrutura",
                "Descrição do erro": "'Examples' encontrado em um 'Scenario' comum."
            })


def ler_feature(caminho_relativo,caminho_arquivo, erros):
    parser = Parser()
    try:
        with open(caminho_arquivo, "r", encoding="utf-8") as arquivo:
            conteudo = arquivo.read()
        data = parser.parse(conteudo)
        # print(data)
        dado = data["feature"]
        validar_feature_estrutura(dado,caminho_relativo, caminho_arquivo, erros)
        # pprint.pprint(dado)
        if 'name' in dado and not dado['name']:
            erros.append({"Arquivo": caminho_relativo,
                          "Tipo do erro": "Feature",
                          "Descrição do erro": "Feature sem título."})
        if 'description' in dado:
            desc = dado["description"]
            padrao_incorreto = re.compile(r'\bBackground(?!:)\b|\bBackground\s+(?!:)\b')
            if padrao_incorreto.search(desc):
                erros.append({"Arquivo": caminho_relativo,
                              "Tipo do erro": "Syntax",
                              "Descrição do erro": "'Background' encontrado sem ':' ou com erro de escrita."})
        tem_background = any('background' in item for item in dado.get('children', []))
        kyw_count_bck = {'Given': 0}
        if tem_background:
            for item in dado.get('children', []):
                if 'background' in item:
                    bck = item['background']
                    for stp in bck.get('steps', []):
                        if stp['keyword'].strip() == 'Given':
                            kyw_count_bck['Given'] += 1
            if kyw_count_bck['Given'] == 0:
                erros.append({"Arquivo": caminho_relativo,
                              "Tipo do erro": "Background",
                              "Descrição do erro": "O background não contém nenhum 'Given'."})
            if kyw_count_bck['Given'] > 1:
                erros.append({"Arquivo": caminho_relativo,
                              "Tipo do erro": "Background",
                              "Descrição do erro": f"A keyword 'Given' aparece {kyw_count_bck['Given']} vezes no background."})
        background_tem_given = tem_background and kyw_count_bck['Given'] > 0
        for item in dado.get('children', []):
            if 'scenario' in item:
                scn = item['scenario']
                kyw_count_scn = {'Given': 0, 'When': 0, 'Then': 0}
                for stp in scn.get('steps', []):
                    kw = stp['keyword'].strip()
                    if kw in kyw_count_scn:
                        kyw_count_scn[kw] += 1
                if kyw_count_scn['When'] == 0:
                    erros.append({"Arquivo": caminho_relativo,
                                  "Tipo do erro": "Scenario",
                                  "Descrição do erro": f"Cenário '{scn['name']}' não contém 'When'."})
                if kyw_count_scn['Then'] == 0:
                    erros.append({"Arquivo": caminho_relativo,
                                  "Tipo do erro": "Scenario",
                                  "Descrição do erro": f"Cenário '{scn['name']}' não contém 'Then'."})
                for kw, cnt in kyw_count_scn.items():
                    if cnt > 1:
                        erros.append({"Arquivo": caminho_relativo,
                                      "Tipo do erro": "Scenario",
                                      "Descrição do erro": f"Keyword '{kw}' aparece {cnt} vezes no cenário '{scn['name']}'."})
                if background_tem_given and kyw_count_scn['Given'] > 0:
                    erros.append({"Arquivo": caminho_relativo,
                                  "Tipo do erro": "Scenario",
                                  "Descrição do erro": f"Cenário '{scn['name']}' não pode conter 'Given' porque existe 'Given' no Background."})
                steps = scn.get('steps', [])
                if background_tem_given:
                    ordemErrada, linha = validar_ordem_keywords_semGiven(steps)
                    if not ordemErrada:
                        erros.append({"Arquivo": caminho_relativo,
                                      "Tipo do erro": "Ordem",
                                      "Descrição do erro": f"Ordem incorreta no cenário '{scn['name']}'. Palavra-chave na linha '{linha}' com inconsistência."})
                else:
                    ordem_correta, linha = validar_ordem_keywords(steps)
                    if not ordem_correta:
                        erros.append({"Arquivo": caminho_relativo,
                                    "Tipo do erro": "Ordem",
                                    "Descrição do erro": f"Ordem incorreta no cenário '{scn['name']}'. Palavra-chave na linha '{linha}' com inconsistência."})

        if tem_background:
            for item in dado.get('children', []):
                if 'background' in item:
                    for stp in item['background'].get('steps', []):
                        kyw = stp['keyword'].strip()
                        if kyw in ['When', 'Then', 'But']:
                            erros.append({"Arquivo": caminho_relativo,
                                          "Tipo do erro": "Background",
                                          "Descrição do erro": f"Palavra-chave inválida no Background: '{kyw}'."})
        gerar_json_analise(data)
    except Exception as e:
        handler = FeatureErrorHandler(str(e))
        erros.append({"Arquivo": caminho_relativo,
                      "Tipo do erro": "Parsing",
                      "Descrição do erro": handler.parse()})


def validar_ordem_keywords(steps):
    ordem = {'Given': 1, 'When': 2, 'Then': 3}
    ultima = 0
    for stp in steps:
        k = stp['keyword'].strip()
        if k in ordem:
            if ordem[k] < ultima:
                return False, stp.get('location', {}).get('line', None)
            ultima = ordem[k]
    return True, None


def validar_ordem_keywords_semGiven(steps):
    ordem = {'When': 1, 'Then': 2}
    ultima = 0 #2
    for stp in steps:
        k = stp['keyword'].strip()
        if k in ordem:
            if ordem[k] < ultima:
                return False, stp.get('location', {}).get('line', None)
            ultima = ordem[k]
    return True, None

def gerar_relatorios(erros, nome_projeto):
    df = pd.DataFrame(erros)
    excel_path = f"{nome_projeto}_erros_features.xlsx"
    word_path = f"{nome_projeto}_erros_features.docx"
    with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Erros')
        ws = writer.book['Erros']
        for col_cells in ws.columns:
            max_len = 0
            for cell in col_cells:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            col_letter = get_column_letter(col_cells[0].column)
            ws.column_dimensions[col_letter].width = max_len + 2
        ws.auto_filter.ref = ws.dimensions
    doc = Document()
    doc.add_heading(f'Relatório de Erros - {nome_projeto}', level=1)
    for err in erros:
        doc.add_paragraph(f"Arquivo: {err['Arquivo']}")
        doc.add_paragraph(f"Tipo do erro: {err['Tipo do erro']}")
        doc.add_paragraph(f"Descrição do erro: {err['Descrição do erro']}")
        doc.add_paragraph('---')
    doc.save(word_path)
    print(f"Relatórios gerados: {excel_path} e {word_path}")



def validar_diretorio(diretorio, erros):
    print(f"Validando diretório: {diretorio}")
    for root, dirs, files in os.walk(diretorio):
        for file in files:
            caminho_arquivo = os.path.join(root, file)
            caminho_arquivo = os.path.normpath(caminho_arquivo)  # Normaliza o path para o SO atual
            _, _, caminho_relativo = caminho_arquivo.partition(os.sep + 'QA')  # Usa separador certo
            caminho_relativo = 'QA' + caminho_relativo if caminho_relativo else caminho_arquivo
            if file.endswith('.feature'):
                ler_feature(caminho_relativo, caminho_arquivo, erros)
            else:
                erros.append({
                    "Arquivo": caminho_relativo,
                    "Tipo do erro": "Extensão",
                    "Descrição do erro": "Extensão inválida."
                })


# Função para gerar o JSON com dados derivados
def gerar_json_analise(feature_json):
    if isinstance(feature_json, dict):
        feature_json = [feature_json]

    resultado = []

    for feature in feature_json:
        if not isinstance(feature, dict):
            continue

        feature_data = feature.get('feature')
        if not isinstance(feature_data, dict):
            continue

        nome_feature = feature_data.get('name', 'Unknown Feature')
        tags_feature = [tag.get('name') for tag in feature_data.get('tags', []) if isinstance(tag, dict)]

        children = feature_data.get('children', [])
        lista_cenarios = []

        automatizados_count = 0
        execucao_automatizado_count = 0

        for child in children:
            cenario_data = child.get('scenario')
            if not cenario_data:
                continue

            nome_cenario = cenario_data.get('name', 'Unknown Scenario')
            tipo = cenario_data.get('keyword', 'Scenario')
            tags_cenario = [tag.get('name') for tag in cenario_data.get('tags', []) if isinstance(tag, dict)]

            qtd_execucoes = 1
            if tipo.lower() == 'scenario outline':
                examples = cenario_data.get('examples', [])
                if examples and isinstance(examples, list):
                    table_body = examples[0].get('tableBody', [])
                    if isinstance(table_body, list):
                        qtd_execucoes = len(table_body)

            if '@automatizado' in tags_cenario:
                automatizados_count += 1
                execucao_automatizado_count += qtd_execucoes

            lista_cenarios.append({
                'cenario': nome_cenario,
                'tipo': tipo,
                'tags': tags_cenario,
                'qtd_execucoes': qtd_execucoes
            })

        resultado.append({
            'feature': nome_feature,
            'tags': tags_feature,
            'automatizados': automatizados_count,
            'execucao_automatizado': execucao_automatizado_count,
            'cenarios': lista_cenarios
        })

    # Verifica se o arquivo já existe e carrega o conteúdo anterior
    if os.path.exists("features.json"):
        with open("features.json", "r", encoding="utf-8") as f:
            conteudo_anterior = json.load(f)
    else:
        conteudo_anterior = []

    # Junta os resultados antigos com os novos
    conteudo_atualizado = conteudo_anterior + resultado

    # Reescreve o arquivo com tudo
    with open("features.json", "w", encoding="utf-8") as f:
        json.dump(conteudo_atualizado, f, ensure_ascii=False, indent=4)

    return resultado


def iniciar_validacao(caminho, nome_projeto):
    print(f"Iniciando validação para o projeto: {nome_projeto}")
    # print(f"Pasta de features: {caminho}")
    
    # Remover arquivo features.json se existir
    if os.path.exists("features.json"):
        os.remove("features.json")
        print("Arquivo features.json removido.")
    
    nome_projeto = nome_projeto.strip().replace(" ", "_")
    if not nome_projeto:
        print("ERRO: Nome do projeto não informado.")
        return False
    
    if not os.path.isdir(caminho):
        print(f"ERRO: Pasta '{caminho}' não existe.")
        return False
    
    erros = []
    validar_diretorio(caminho, erros)
    
    gerar_relatorios(erros, nome_projeto)
    
    print(f"Total de erros encontrados: {len(erros)}")
    # for i, err in enumerate(erros, 1):
    #     print(f"{i}. Arquivo: {err['arquivo']}")
    #     print(f"   Tipo: {err['tipo']}")
    #     print(f"   Mensagem: {err['mensagem']}")
    #     print("---")
    
    # Gerar planilha de regressão
    try:
        json_file = "features.json"
        with open(json_file, "r", encoding="utf-8") as f:
            json_data = json.load(f)
        
        # Chamando no metodo tradicional de importação
        try:
            from planilha_regressao_json_jenkins import gerar_planilha_regressao, converte_json
            output_file = gerar_planilha_regressao(converte_json(json_data), nome_projeto)
            print(f"Planilha de regressão gerada: {output_file}")
        except ImportError:
            # Caso dê erro no processo padrao, chamo como subprocesso
            import subprocess
            cmd = [sys.executable, "planilha_regressao_json_jenkins.py", json_file, nome_projeto]
            print(f"Executando comando: {' '.join(cmd)}")
            result = subprocess.run(cmd, capture_output=True, text=True)
            
            if result.returncode != 0:
                print(f"ERRO ao executar planilha_regressao_json_jenkins.py: {result.stderr}")
                return False
            
            print(result.stdout)
            
        print(f"Planilha de regressão gerada com sucesso: Regressao_{nome_projeto}.xlsx")
    except Exception as e:
        print(f"ERRO ao gerar planilha de regressão: {str(e)}")
        return False
    
    print("Validação concluída com sucesso!")
    return True

def main():
    # Verificar se os argumentos foram passados
    if len(sys.argv) == 3:
        # via argumentos de linha de comando
        caminho = sys.argv[1]
        nome_projeto = sys.argv[2]
        iniciar_validacao(caminho, nome_projeto)
    else:
        # parâmetros via variáveis de ambiente
        caminho = os.environ.get('CAMINHO_PASTA_FEATURES')
        nome_projeto = os.environ.get('NOME_PROJETO_ANALISADO')
        
        if caminho and nome_projeto:
            iniciar_validacao(caminho, nome_projeto)
        else:
            print("ERRO: Parâmetros insuficientes.")
            print("Uso: python validador_features_headless.py <caminho_pasta_features> <nome_projeto>")
            print("Ou defina as variáveis de ambiente CAMINHO_PASTA_FEATURES e NOME_PROJETO_ANALISADO")
            sys.exit(1)

if __name__ == "__main__":
    main()
